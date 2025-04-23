import openai
import streamlit as st
import os
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from docx import Document
from dotenv import load_dotenv
from google.oauth2.service_account import Credentials

# Charger les variables d'environnement depuis le fichier .env
load_dotenv()

# Récupérer la clé API OpenAI depuis le fichier .env
API_KEY = os.getenv("OPENAI_API_KEY")

# Vérifier si la clé API a été correctement chargée
if not API_KEY:
    st.error("La clé API OpenAI est manquante. Assurez-vous qu'elle est dans le fichier .env.")
else:
    openai.api_key = API_KEY

# --- Fonction pour récupérer les données de Google Sheets ---
# Chemin du fichier .json et ID du dossier Google Drive
SERVICE_ACCOUNT_FILE = '/Users/ahmeddiomande/Documents/IDEALMATCH2025/generateur-fiches-clean/.devcontainer/peak-dominion-453716-v3-502cdb22fa02.json'
FOLDER_ID = '19uTtwgjpK4tYjcENCSDbpLCtAgK_sGkW'  # ID de ton dossier Google Drive

# Créer les identifiants d'authentification pour l'API Google Sheets
credentials = Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=["https://www.googleapis.com/auth/spreadsheets.readonly"])

# Construire le service pour l'API Google Sheets
service = build('sheets', 'v4', credentials=credentials)

SPREADSHEET_ID = '1wl_OvLv7c8iN8Z40Xutu7CyrN9rTIQeKgpkDJFtyKIU'  # ID de ton fichier Google Sheets
RANGE_NAME = 'Besoins ASI!A1:Z1000'  # Plage des données

# Fonction pour récupérer les données depuis la feuille de calcul
def recuperer_donnees_google_sheet():
    sheet = service.spreadsheets()
    result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=RANGE_NAME).execute()
    values = result.get('values', [])
    return values

# --- Fonction pour générer la fiche de poste dans Google Drive ---
def enregistrer_fiche_google_drive(nom_fichier, contenu_fiche):
    doc = Document()
    doc.add_heading(f'Fiche de Poste : {nom_fichier}', 0)
    doc.add_paragraph(contenu_fiche)
    
    # Sauvegarder le fichier localement
    chemin_fichier = f'/content/drive/My Drive/{nom_fichier}.docx'
    doc.save(chemin_fichier)

    # Ajouter le fichier à Google Drive
    file_metadata = {
        'name': f'{nom_fichier}.docx',
        'parents': [FOLDER_ID]
    }

    media = MediaFileUpload(chemin_fichier, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    drive_service = build('drive', 'v3', credentials=credentials)
    
    file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()

    return f"Fichier généré et stocké sur Google Drive : https://drive.google.com/file/d/{file['id']}/view"

# --- Interface utilisateur Streamlit ---
st.title('🎯 IDEALMATCH JOB CREATOR')

# Texte introductif
st.markdown("""
Bienvenue dans l'outil **IDEALMATCH JOB CREATOR** !  
Cet outil vous permet de générer des fiches de poste personnalisées à l'aide de l'intelligence artificielle (ChatGPT).

### Instructions :
- Entrez votre **prompt personnalisé** dans la zone de texte ci-dessous.
- Cliquez sur le bouton "Générer la Fiche de Poste" pour obtenir une fiche automatiquement générée.
- La fiche sera basée sur votre description du poste et des critères de sélection.

📝 **Astuces** :
- Soyez précis dans votre description pour obtenir les meilleurs résultats.
- L'outil utilise la dernière version de GPT-3.5 pour vous fournir des résultats de qualité.
""")

# --- Zone de saisie du prompt de l'utilisateur ---
user_prompt = st.text_area("Écrivez ici votre prompt pour générer une fiche de poste :", 
                          "Entrez ici le prompt pour ChatGPT...")

# --- Bouton pour envoyer la demande à OpenAI ---
if st.button('Générer la Fiche de Poste'):
    if user_prompt:
        try:
            # Appeler l'API OpenAI avec le prompt de l'utilisateur
            response = openai.Completion.create(
                model="gpt-3.5-turbo",  # Ou gpt-4 si tu l'as
                prompt=user_prompt,
                max_tokens=500
            )
            
            # Afficher la réponse générée par ChatGPT
            st.subheader('Fiche de Poste Générée:')
            st.write(response.choices[0].text.strip())
        
        except Exception as e:
            st.error(f"Erreur lors de la génération de la fiche de poste : {e}")
    else:
        st.warning("Veuillez entrer un prompt avant de soumettre.")

# --- Bouton pour générer la fiche de poste à partir de Google Sheets ---
if st.button('Générer à partir du fichier RPO'):
    # Récupérer les données du fichier RPO
    donnees_rpo = recuperer_donnees_google_sheet()

    for ligne in donnees_rpo[1:]:
        poste_selectionne = dict(zip(donnees_rpo[0], ligne))  # Associer les colonnes aux valeurs de la ligne

        # Créer le prompt pour OpenAI avec les données récupérées
        prompt_rpo = f"""Générer une fiche de poste pour le poste {poste_selectionne['Titre du poste recherché']} en utilisant les informations suivantes :
        - Durée de la mission : {poste_selectionne['Durée de la mission']}
        - Statut mission : {poste_selectionne['statut mission']}
        - Compétences obligatoires : {poste_selectionne['Compétences obligatoires ( Préciser technologies principales et frameworks pour les postes techniques )']}
        - Projet : {poste_selectionne['Projet sur lequel va travailler le ou la candidate :']}
        - Localisation : {poste_selectionne['Localisation']}
        """

        try:
            # Appeler l'API OpenAI avec le prompt généré
            response_rpo = openai.Completion.create(
                model="gpt-3.5-turbo",  # Ou gpt-4 si tu l'as
                prompt=prompt_rpo,
                max_tokens=500
            )

            # Sauvegarder la fiche de poste dans Google Drive
            contenu_fiche = response_rpo.choices[0].text.strip()
            nom_fichier = f"Fiche_{poste_selectionne['Titre du poste recherché']}"
            lien_drive = enregistrer_fiche_google_drive(nom_fichier, contenu_fiche)

            st.subheader(f"Fiche de Poste générée pour : {poste_selectionne['Titre du poste recherché']}")
            st.write(contenu_fiche)
            st.write(f"Fiche stockée sur Google Drive : {lien_drive}")
        
        except Exception as e:
            st.error(f"Erreur lors de la génération de la fiche de poste pour {poste_selectionne['Titre du poste recherché']}: {e}")
