import openpyxl
import requests
from docx import Document
import os
import zipfile
from googleapiclient.discovery import build
from google.oauth2.service_account import Credentials

# Clé API OpenAI
from dotenv import load_dotenv
load_dotenv()
API_KEY = os.getenv("OPENAI_API_KEY")

# Nettoyer la stack technique (séparer par virgules)
def formater_stack(stack_brut):
    if not stack_brut:
        return "À définir"
    séparateurs = [",", "/", "|", ";", "\n"]
    for sep in séparateurs:
        if sep in stack_brut:
            éléments = [elem.strip() for elem in stack_brut.split(sep) if elem.strip()]
            return ", ".join(éléments)
    return stack_brut.strip()

# Fonction pour générer la fiche de poste
def generer_fiche_poste(data):
    stack_formatee = formater_stack(data["Compétences obligatoires"])
    prompt = f"""
Présentation de la société :
Société fondée en [Année] spécialisée dans le secteur {data["Secteur d'activité"]}. Elle commercialise des solutions [Produit/Solution] à destination de [Clients cibles]. Elle compte actuellement [Nombre de salariés] collaborateurs, localisée à {data["Localisation"]}.

Description du poste :
Vous rejoignez une équipe de {data["Taille de l'équipe"]} personnes sur un projet dans le domaine {detecter_domaine(stack_formatee)}.
Nous recherchons un(e) {data["Titre du poste recherché"]} pour renforcer l'équipe dans un contexte de {data["Durée de la mission"]}.

Missions :
{data["Projet"] if data["Projet"] else generer_missions_generiques(data["Titre du poste recherché"])}
- Participer activement aux développements.
- Assurer la qualité du livrable.
- Collaborer avec les équipes métiers et techniques.

Stack technique :
{stack_formatee}

Profil recherché :
- {data["Nombre d'années d'expérience"]} ans minimum d'expérience.
- Maîtrise de : {stack_formatee}.
- Capacité à travailler en équipe ({data["Taille de l'équipe"]} personnes).
- Expérience sur des projets similaires souhaitée.

Avantages :
{"TJM : " + data["TJM"] if data["Statut"].lower() == "freelance" else "Salaire : " + (data["Salaire"] if data["Salaire"] else "Salaire à déterminer selon profil")}
Télétravail : {data["Télétravail"]}
Date de démarrage : {data["Date de démarrage"]}

Rédige une fiche de poste professionnelle en suivant strictement cette structure.
"""

    headers = {"Authorization": f"Bearer {API_KEY}"}
    payload = {
        "model": "gpt-3.5-turbo",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3
    }

    response = requests.post(
        "https://api.openai.com/v1/chat/completions",
        json=payload,
        headers=headers
    )

    if response.status_code != 200:
        print(f"Erreur API : {response.status_code}")
        return "Erreur lors de la génération."

    result = response.json()
    return result["choices"][0]["message"]["content"]

def generer_mail_type(data):
    return (
        f"Bonjour [Nom du Candidat],\n\n"
        f"Je vous contacte suite à votre profil qui a retenu mon attention.\n"
        f"Nous avons actuellement une opportunité intéressante : {data['Titre du poste recherché']} chez {data['Nom du client']}.\n"
        f"Seriez-vous disponible pour en discuter ?\n\n"
        f"Bien cordialement,\n[Votre Nom]"
    )

# Détection domaine simplifié
def detecter_domaine(stack):
    if any(keyword in stack.lower() for keyword in ["data", "big data", "python", "sql"]):
        return "Data & IT"
    elif any(keyword in stack.lower() for keyword in ["finance", "banque", "assurance"]):
        return "Finance"
    elif any(keyword in stack.lower() for keyword in ["industrie", "production"]):
        return "Industrie"
    return "Tech"

def generer_missions_generiques(titre_poste):
    return f"- Contribuer aux missions liées au poste de {titre_poste}.\n- Assurer un suivi rigoureux des livrables."

# --- Authentification Google Sheets ---
SERVICE_ACCOUNT_FILE = '/path/to/your/service/account/file.json'
SPREADSHEET_ID = '1wl_OvLv7c8iN8Z40Xutu7CyrN9rTIQeKgpkDJFtyKIU'
RANGE_NAME = 'Besoins ASI!A1:Z1000'
EXCEL_FILE_PATH = 'data/RPO.xlsx'

credentials = Credentials.from_service_account_file(SERVICE_ACCOUNT_FILE, scopes=["https://www.googleapis.com/auth/spreadsheets.readonly"])
service = build('sheets', 'v4', credentials=credentials)

# Supprimer ancien fichier
if os.path.exists(EXCEL_FILE_PATH):
    os.remove(EXCEL_FILE_PATH)
    print(f"Ancien fichier {EXCEL_FILE_PATH} supprimé avec succès.")

# Télécharger les données
sheet = service.spreadsheets()
result = sheet.values().get(spreadsheetId=SPREADSHEET_ID, range=RANGE_NAME).execute()
values = result.get('values', [])

# Créer le nouveau Excel
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Feuil1"
for row in values:
    ws.append(row)
wb.save(EXCEL_FILE_PATH)
print(f"Nouveau fichier {EXCEL_FILE_PATH} créé avec succès.")

# Nettoyage nom fichier
def nettoyer_nom_fichier(nom):
    return nom.replace("/", "_").replace("\\", "_").replace(":", "_").replace("?", "_").replace("*", "_").replace("\"", "_").replace("<", "_").replace(">", "_").replace("|", "_").replace("\n", "_")

# Charger Excel
wb_rpo = openpyxl.load_workbook(EXCEL_FILE_PATH)
sheet_rpo = wb_rpo["Feuil1"]

# --- Génération des documents ---
fichiers_generes = []  # Liste des fichiers à ajouter au ZIP

for row in sheet_rpo.iter_rows(min_row=2, values_only=True):
    if len(row) < 6:
        continue

    data = {
        "Nom du client": row[4] if len(row) > 4 else "",
        "Secteur d'activité": row[3] if len(row) > 3 else "Non précisé",
        "Localisation": row[5] if len(row) > 5 else "",
        "Titre du poste recherché": row[6] if len(row) > 6 else "",
        "Statut": row[7] if len(row) > 7 else "",
        "Nombre d'années d'expérience": row[8] if len(row) > 8 else "",
        "Date de démarrage": row[10] if len(row) > 10 else "",
        "TJM": row[11] if len(row) > 11 else "",
        "Salaire": row[12] if len(row) > 12 else "",
        "Télétravail": row[13] if len(row) > 13 else "",
        "Durée de la mission": row[14] if len(row) > 14 else "",
        "Projet": row[15] if len(row) > 15 else "",
        "Compétences obligatoires": row[16] if len(row) > 16 else "",
        "Taille de l'équipe": row[17] if len(row) > 17 else ""
    }

    nom_fichier_base = nettoyer_nom_fichier(f"{data['Nom du client']}_{data['Titre du poste recherché']}")
    chemin_fichier_word = f"fiches/{nom_fichier_base}.docx"
    chemin_mail_word = f"mails/{nom_fichier_base}_MAIL.docx"

    # Générer Fiche
    fiche_poste = generer_fiche_poste(data)
    document_fiche = Document()
    document_fiche.add_heading(f"Fiche de Poste : {data['Titre du poste recherché']}", level=1)
    document_fiche.add_paragraph(fiche_poste)
    document_fiche.save(chemin_fichier_word)

    # Générer et sauvegarder le mail type
    mail_type = generer_mail_type(data)
    document_mail = Document()
    document_mail.add_heading(f"Mail Type : {data['Titre du poste recherché']} chez {data['Nom du client']}", level=1)
    document_mail.add_paragraph(mail_type)
    document_mail.save(chemin_mail_word)

    # Ajouter les fichiers générés à la liste
    fichiers_generes.append(chemin_fichier_word)
    fichiers_generes.append(chemin_mail_word)

    print(f"Fiche de poste enregistrée : {chemin_fichier_word}")
    print(f"Mail type enregistré : {chemin_mail_word}")

# Créer un fichier ZIP contenant RPO, les fiches et les mails
def creer_zip(chemin_zip, fichiers):
    with zipfile
