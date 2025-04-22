import pandas as pd
import os
from datetime import datetime
import openai
from docx import Document

# --- CONFIGURATION ---
from dotenv import load_dotenv
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

CSV_PATH = str(Path(__file__).parent / "data" / "Suivi_des_opportunités.csv")
EXCEL_PATH = str(Path(__file__).parent / "data" / "fichier_cible.xlsx")
MAIL_DIR = str(Path(__file__).parent / "redaction_mail")
PROMPT_DIR = str(Path(__file__).parent / "prompts")

DOSSIERS_PROMPTS = {
    "Annonces_anonymes": {
        "colonne": "Annonces_anonymes",
        "chemin": str(Path(__file__).parent / "fiches" / "Annonces_anonymes"),
        "prompt_file": "anonyme.txt"
    },
    "Annonce_ASI_interne": {
        "colonne": "Annonce_ASI_interne",
        "chemin": str(Path(__file__).parent / "fiches" / "Annonce_ASI_interne"),
        "prompt_file": "interne.txt"
    },
    "Annonce_client_ASI": {
        "colonne": "Annonce_client_ASI",
        "chemin": "/Users/ahmeddiomande/Documents/IDEALMATCH2025/JOB CREATOR/Fiche_de_poste/Annonce_client_ASI",
        "prompt_file": "client.txt"
    },
    "Annonce_freelance": {
        "colonne": "Annonce_freelance",
        "chemin": "/Users/ahmeddiomande/Documents/IDEALMATCH2025/JOB CREATOR/Fiche_de_poste/Annonce_freelance",
        "prompt_file": "freelance.txt"
    },
}

# --- Lecture CSV et nettoyage ---
encodings = ['utf-8', 'ISO-8859-1', 'cp1252']
for enc in encodings:
    try:
        df_csv = pd.read_csv(CSV_PATH, encoding=enc)
        print(f"✅ CSV chargé avec encodage : {enc}")
        break
    except UnicodeDecodeError:
        continue
else:
    raise ValueError("❌ Impossible de lire le CSV avec les encodages courants.")

col_mapping = {
    "Ingénieur d'affaires": "Ingénieur",
    "Clients": "Client",
    "Profil": "Profil",
    "Résumé de la mission": "Résumé",
    "Budget / TJM proposé": "TJM",
    "Localisation": "Localisation"
}
df_csv = df_csv[list(col_mapping.keys())].rename(columns=col_mapping)

# --- Nettoyage via ChatGPT ---
def clean_text(text):
    if pd.isna(text) or not str(text).strip():
        return ""
    prompt = f"Nettoie ce texte HTML ou encodé pour obtenir une phrase claire :\n\nTexte brut :\n{text}\n\nTexte nettoyé :"
    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("❌ Erreur GPT:", e)
        return text

df_csv["Profil"] = df_csv["Profil"].apply(clean_text)
df_csv["Résumé"] = df_csv["Résumé"].apply(clean_text)

# --- Intégration dans Excel ---
if os.path.exists(EXCEL_PATH):
    df_excel = pd.read_excel(EXCEL_PATH)
    df_merge = pd.merge(df_csv, df_excel, how='left', on=['Ingénieur', 'Client', 'Profil', 'Résumé', 'TJM'], indicator=True)
    df_new = df_csv[df_merge['_merge'] == 'left_only'].copy()
else:
    df_excel = pd.DataFrame(columns=[*df_csv.columns, 'Numéro de série', 'Date d\'insertion'] + [conf["colonne"] for conf in DOSSIERS_PROMPTS.values()])
    df_new = df_csv.copy()

start_index = 1
if not df_excel.empty:
    ids = df_excel['Numéro de série'].dropna().tolist()
    nums = [int(i[3:]) for i in ids if str(i).startswith("JOB") and str(i[3:]).isdigit()]
    if nums:
        start_index = max(nums) + 1

df_new['Numéro de série'] = [f"JOB{str(i).zfill(4)}" for i in range(start_index, start_index + len(df_new))]
df_new["Date d'insertion"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

# --- Fusionner ---
df_final = pd.concat([df_excel, df_new], ignore_index=True).reset_index(drop=True)
df_final = df_final.dropna(subset=['Profil', 'Client', 'Résumé'], how='all')

# --- Génération des fichiers ---
def nettoyer_nom(n):
    return str(n).replace("/", "_").replace("\\", "_").replace(":", "_").replace("*", "_").replace("?", "_").replace("\"", "_").replace("<", "_").replace(">", "_").replace("|", "_")

def generer_docx(chemin, titre, contenu):
    doc = Document()
    doc.add_heading(titre, level=1)
    doc.add_paragraph(contenu)
    doc.save(chemin)

def generer_contenu(chemin_prompt_txt, row):
    try:
        with open(os.path.join(PROMPT_DIR, chemin_prompt_txt), 'r', encoding='utf-8') as f:
            prompt_template = f.read()

        donnees_ligne = "\n".join([f"{col} : {str(row[col])}" for col in df_final.columns if col not in DOSSIERS_PROMPTS])

        prompt = (
            prompt_template
            .replace("{Profil}", str(row.get("Profil", "")))
            .replace("{Client}", str(row.get("Client", "")))
            .replace("{Résumé}", str(row.get("Résumé", "")))
            .replace("{Localisation}", str(row.get("Localisation", "")))
            .replace("{DONNEES}", donnees_ligne)
        )

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("Erreur de génération GPT:", e)
        return "Erreur de génération."

nouvelles_lignes = 0
lignes_ignorees = 0

for i, row in df_final.iterrows():
    print(f"\n--- Ligne {i+1} ---")
    for conf in DOSSIERS_PROMPTS.values():
        print(f"{conf['colonne']}: {row.get(conf['colonne'])}")

    if all(pd.isna(row.get(conf["colonne"])) for conf in DOSSIERS_PROMPTS.values()):
        print(f"✅ Traitement de la ligne {i+1} ({row.get('Profil')} chez {row.get('Client')})")

        id_ = nettoyer_nom(str(row.get('Numéro de série', f"LIGNE{i+1}")))
        profil = nettoyer_nom(row.get("Profil", "Profil"))
        client = nettoyer_nom(row.get("Client", "Client"))

        for type_annonce, conf in DOSSIERS_PROMPTS.items():
            os.makedirs(conf["chemin"], exist_ok=True)
            nom_fichier = f"{profil}_{client}_{type_annonce}.docx"
            chemin_complet = os.path.join(conf["chemin"], nom_fichier)

            contenu = generer_contenu(conf["prompt_file"], row)
            generer_docx(chemin_complet, f"Fiche de poste - {profil}", contenu)

            formule = f'=HYPERLINK("file://{chemin_complet.replace(" ", "%20")}", "📄 Voir l\'annonce")'
            df_final.at[i, conf["colonne"]] = str(formule)

        os.makedirs(MAIL_DIR, exist_ok=True)
        mail_path = os.path.join(MAIL_DIR, f"{profil}_{client}_MAIL.docx")
        generer_docx(mail_path, f"Mail {profil}", f"Bonjour [Candidat],\n\nNous avons une opportunité {profil} chez {client}. Intéressé ?\n\nBien cordialement.")

        nouvelles_lignes += 1
    else:
        print(f"⏭ Ligne {i+1} ignorée (fiches déjà présentes)")
        lignes_ignorees += 1

# --- Sauvegarder ---
try:
    df_final.to_excel(EXCEL_PATH, index=False)
    print("💾 Fichier Excel mis à jour avec succès.")
except PermissionError:
    print("❌ Permission refusée : ferme le fichier Excel avant de lancer le script.")

print("\n--- RÉSUMÉ ---")
print(f"✅ {nouvelles_lignes} ligne(s) traitée(s) avec fiches et mails.")
print(f"⏭ {lignes_ignorees} ligne(s) déjà complétées ignorées.")
print("📁 Tous les fichiers ont été générés avec succès.")
